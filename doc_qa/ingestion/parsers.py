import csv
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Module-level imports so tests can patch them via "doc_qa.ingestion.parsers.*"
try:
    import pytesseract
    from pdf2image import convert_from_path

    _OCR_AVAILABLE = True
except ImportError:
    pytesseract = None  # type: ignore[assignment]
    convert_from_path = None  # type: ignore[assignment]
    _OCR_AVAILABLE = False

ParsedPage = tuple[str, int]

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".json", ".docx"}


def parse_pdf(path: Path) -> list[ParsedPage]:
    import pypdf

    reader = pypdf.PdfReader(str(path))
    pages: list[ParsedPage] = []

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append((text, i))

    # Decide whether OCR fallback is needed
    from config.settings import settings

    mean_chars = sum(len(t) for t, _ in pages) / len(pages) if pages else 0.0

    if mean_chars < settings.pdf_ocr_threshold:
        pages = _ocr_pdf(path, reader)

    return [(t, n) for t, n in pages if t.strip()]


def _ocr_pdf(path: Path, reader: object) -> list[ParsedPage]:
    if not _OCR_AVAILABLE or convert_from_path is None or pytesseract is None:
        logger.error(
            "PDF %s: OCR fallback required but pytesseract/pdf2image not installed. "
            "Returning pypdf text.",
            path.name,
        )
        return _pypdf_pages(path)

    logger.warning("PDF %s: OCR fallback triggered", path.name)
    try:
        images = convert_from_path(str(path))
    except Exception as exc:
        logger.error("PDF %s: pdf2image failed: %s. Returning empty.", path.name, exc)
        return []

    pages: list[ParsedPage] = []
    for i, image in enumerate(images, start=1):
        try:
            text = pytesseract.image_to_string(image).strip()
        except pytesseract.TesseractNotFoundError:
            logger.error("PDF %s: Tesseract binary not found. Returning pypdf text.", path.name)
            return _pypdf_pages(path)
        if text:
            pages.append((text, i))
    return pages


def _pypdf_pages(path: Path) -> list[ParsedPage]:
    import pypdf

    pages = []
    for i, page in enumerate(pypdf.PdfReader(str(path)).pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((text, i))
    return pages


def parse_txt(path: Path) -> list[ParsedPage]:
    try:
        lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    except OSError:
        return []
    pages: list[ParsedPage] = []
    for i, line in enumerate(lines, start=1):
        text = line.strip()
        if text:
            pages.append((text, i))
    return pages


def parse_markdown(path: Path) -> list[ParsedPage]:
    # Same as parse_txt but heading markers are preserved
    return parse_txt(path)


def parse_csv(path: Path) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return []
    reader = csv.DictReader(text.splitlines())
    for i, row in enumerate(reader, start=1):
        serialized = json.dumps(dict(row)).strip()
        if serialized:
            pages.append((serialized, i))
    return pages


def parse_json(path: Path) -> list[ParsedPage]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return []
    try:
        data: list | dict | str | int | float = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON in file: {path}") from exc

    pages: list[ParsedPage] = []
    if isinstance(data, list):
        for i, element in enumerate(data, start=1):
            serialized = json.dumps(element).strip()
            if serialized:
                pages.append((serialized, i))
    elif isinstance(data, dict):
        for i, (key, value) in enumerate(data.items(), start=1):
            serialized = f"{key}: {json.dumps(value)}".strip()
            if serialized:
                pages.append((serialized, i))
    else:
        serialized = json.dumps(data).strip()
        if serialized:
            pages.append((serialized, 1))
    return pages


def parse_docx(path: Path) -> list[ParsedPage]:
    try:
        import docx
    except ImportError:
        logger.error("python-docx not installed; cannot parse %s", path.name)
        return []

    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        logger.error("Failed to open %s: %s", path.name, exc)
        return []

    pages: list[ParsedPage] = []
    index = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pages.append((text, index))
            index += 1

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells).strip()
            if row_text:
                pages.append((row_text, index))
                index += 1

    return pages


def parse_file(path: Path) -> list[ParsedPage]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file extension '{suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    parsers = {
        ".pdf": parse_pdf,
        ".txt": parse_txt,
        ".md": parse_markdown,
        ".csv": parse_csv,
        ".json": parse_json,
        ".docx": parse_docx,
    }
    return parsers[suffix](path)
