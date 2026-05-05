import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from doc_qa.ingestion.parsers import (
    parse_csv,
    parse_file,
    parse_json,
    parse_markdown,
    parse_pdf,
    parse_txt,
)

FIXTURE_DIR = Path(__file__).parent / "fixtures"


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def test_parse_txt_returns_lines(tmp_path):
    f = tmp_path / "sample.txt"
    f.write_text("Line one\nLine two\nLine three\n")
    result = parse_txt(f)
    assert len(result) == 3
    texts = [t for t, _ in result]
    assert texts == ["Line one", "Line two", "Line three"]
    line_nums = [n for _, n in result]
    assert line_nums == [1, 2, 3]


def test_parse_txt_empty_file(tmp_path):
    f = tmp_path / "empty.txt"
    f.write_text("")
    assert parse_txt(f) == []


def test_parse_txt_skips_blank_lines(tmp_path):
    f = tmp_path / "blanks.txt"
    f.write_text("First\n\n\nFourth\n")
    result = parse_txt(f)
    assert len(result) == 2
    assert result[0][0] == "First"
    assert result[1][0] == "Fourth"


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def test_parse_markdown_preserves_headings(tmp_path):
    f = tmp_path / "doc.md"
    f.write_text("# Heading One\n## Heading Two\nSome body text.\n")
    result = parse_markdown(f)
    texts = [t for t, _ in result]
    assert any(t.startswith("# ") for t in texts)
    assert any(t.startswith("## ") for t in texts)


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------


def test_parse_csv_serializes_rows(tmp_path):
    f = tmp_path / "rates.csv"
    f.write_text("product,rate\n30yr_fixed,6.75\n15yr_fixed,6.25\n")
    result = parse_csv(f)
    assert len(result) == 2
    row1 = json.loads(result[0][0])
    assert row1 == {"product": "30yr_fixed", "rate": "6.75"}
    assert result[0][1] == 1
    assert result[1][1] == 2


def test_parse_csv_empty_returns_empty(tmp_path):
    f = tmp_path / "empty.csv"
    f.write_text("product,rate\n")
    assert parse_csv(f) == []


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------


def test_parse_json_list_root(tmp_path):
    data = [{"id": 1, "name": "Alice"}, {"id": 2, "name": "Bob"}, {"id": 3, "name": "Carol"}]
    f = tmp_path / "data.json"
    f.write_text(json.dumps(data))
    result = parse_json(f)
    assert len(result) == 3
    assert result[0][1] == 1
    assert json.loads(result[0][0]) == data[0]


def test_parse_json_dict_root(tmp_path):
    data = {"key_a": "value_a", "key_b": [1, 2, 3]}
    f = tmp_path / "data.json"
    f.write_text(json.dumps(data))
    result = parse_json(f)
    assert len(result) == 2
    assert result[0][0].startswith("key_a:")
    assert result[1][0].startswith("key_b:")


def test_parse_json_empty_list(tmp_path):
    f = tmp_path / "empty.json"
    f.write_text("[]")
    assert parse_json(f) == []


def test_parse_json_invalid_raises_value_error(tmp_path):
    f = tmp_path / "bad.json"
    f.write_text("{not valid json")
    with pytest.raises(ValueError, match="Invalid JSON"):
        parse_json(f)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_parse_docx_paragraphs(tmp_path):
    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    from doc_qa.ingestion.parsers import parse_docx

    result = parse_docx(path)
    texts = [t for t, _ in result]
    assert "First paragraph." in texts
    assert "Second paragraph." in texts
    assert "Third paragraph." in texts


def test_parse_docx_skips_empty_paragraphs(tmp_path):
    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_paragraph("Real content.")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("More content.")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    from doc_qa.ingestion.parsers import parse_docx

    result = parse_docx(path)
    assert len(result) == 2


def test_parse_docx_includes_table_rows(tmp_path):
    import docx as docx_lib

    doc = docx_lib.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Product"
    table.cell(0, 1).text = "Rate"
    table.cell(1, 0).text = "30yr fixed"
    table.cell(1, 1).text = "6.75%"
    path = tmp_path / "test.docx"
    doc.save(str(path))

    from doc_qa.ingestion.parsers import parse_docx

    result = parse_docx(path)
    all_text = " ".join(t for t, _ in result)
    assert "Product" in all_text or "Rate" in all_text
    assert len(result) >= 2


# ---------------------------------------------------------------------------
# parse_file dispatch
# ---------------------------------------------------------------------------


def test_parse_file_raises_on_unsupported(tmp_path):
    f = tmp_path / "spreadsheet.xlsx"
    f.write_bytes(b"fake xlsx content")
    with pytest.raises(ValueError, match="Unsupported"):
        parse_file(f)


def test_parse_file_dispatches_txt(tmp_path):
    f = tmp_path / "doc.txt"
    f.write_text("Hello world.\n")
    result = parse_file(f)
    assert len(result) == 1
    assert result[0][0] == "Hello world."


# ---------------------------------------------------------------------------
# PDF — pypdf (digital) path
# ---------------------------------------------------------------------------


def test_parse_pdf_pypdf_path(tmp_path):
    """Digital PDF with sufficient text must NOT trigger OCR."""
    import io

    from reportlab.pdfgen import canvas as rl_canvas

    # Build a 2-page PDF with 200+ chars per page
    page_text_1 = "A" * 210
    page_text_2 = "B" * 210
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(72, 720, page_text_1)
    c.showPage()
    c.drawString(72, 720, page_text_2)
    c.save()
    pdf_path = tmp_path / "digital.pdf"
    pdf_path.write_bytes(buf.getvalue())

    with (
        patch("doc_qa.ingestion.parsers.pytesseract", create=True) as mock_tess,
        patch("doc_qa.ingestion.parsers.convert_from_path", create=True) as mock_conv,
    ):
        from doc_qa.ingestion.parsers import parse_pdf

        result = parse_pdf(pdf_path)

    assert len(result) == 2
    mock_conv.assert_not_called()
    mock_tess.assert_not_called() if hasattr(mock_tess, "assert_not_called") else None


# ---------------------------------------------------------------------------
# PDF — OCR fallback
# ---------------------------------------------------------------------------


def test_parse_pdf_ocr_fallback_triggered():
    """Image-only PDF must trigger OCR (uses pre-generated fixture)."""
    fixture = FIXTURE_DIR / "scanned_promissory_note.pdf"
    if not fixture.exists():
        pytest.skip("OCR fixture not found — run tests/fixtures/create_scanned_fixture.py")

    import pytesseract

    from doc_qa.ingestion.parsers import parse_pdf

    try:
        pytesseract.get_tesseract_version()
    except pytesseract.TesseractNotFoundError:
        pytest.skip("Tesseract not installed on this system")

    with (
        patch("doc_qa.ingestion.parsers.convert_from_path") as mock_conv,
        patch("doc_qa.ingestion.parsers.pytesseract") as mock_tess,
    ):
        mock_conv.return_value = [MagicMock(), MagicMock()]  # 2 fake PIL images
        mock_tess.image_to_string.return_value = "Extracted OCR text from scanned page."
        mock_tess.TesseractNotFoundError = pytesseract.TesseractNotFoundError

        result = parse_pdf(fixture)

    mock_conv.assert_called_once()
    assert mock_tess.image_to_string.call_count == 2
    assert len(result) == 2


def test_parse_pdf_tesseract_missing_degrades_gracefully(tmp_path):
    """If Tesseract is missing, parse_pdf must return pypdf results without raising."""
    import io

    import pytesseract
    from reportlab.pdfgen import canvas as rl_canvas

    # Sparse PDF — only 10 chars per page, below the OCR threshold
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(72, 720, "Hi")
    c.showPage()
    c.save()
    pdf_path = tmp_path / "sparse.pdf"
    pdf_path.write_bytes(buf.getvalue())

    with (
        patch("doc_qa.ingestion.parsers.convert_from_path") as mock_conv,
        patch("doc_qa.ingestion.parsers.pytesseract") as mock_tess,
    ):
        mock_conv.return_value = [MagicMock()]
        mock_tess.TesseractNotFoundError = pytesseract.TesseractNotFoundError
        mock_tess.image_to_string.side_effect = pytesseract.TesseractNotFoundError

        result = parse_pdf(pdf_path)

    # Should not raise; returns whatever pypdf extracted
    assert isinstance(result, list)
